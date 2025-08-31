import os
import re
import uuid
import json
import shutil
import pdfplumber
import pandas as pd
from pathlib import Path
from typing import List, Optional
from fastapi import FastAPI, UploadFile, File, HTTPException, status
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from pdf2image import convert_from_path
import pytesseract

# --- Load configuration
from config import settings

app = FastAPI(title="Document(SoF & CP) Processor Backend")

# --- CORS
origins = ["http://localhost:3000", "http://localhost:5173"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Startup
@app.on_event("startup")
def on_startup():
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    print(f"Upload directory '{settings.UPLOAD_DIR}' ensured.")


# ---------------- File Extractors ---------------- #
def pdf_extract(path: Path) -> tuple[List[str], str]:
    """
    Extract text lines from PDF using pdfplumber, with OCR fallback.
    Includes normalization and line merging like standalone pdf_to_text.
    """
    text = ""
    method = "plumber"

    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text() or ""
                text += extracted + "\n"
    except Exception:
        text = ""

    # --- Step 2: OCR fallback if nothing extracted ---
    if not text.strip():
        print("[⚠] Falling back to OCR...")
        method = "ocr"
        pages = convert_from_path(str(path))
        for img in pages:
            text += pytesseract.image_to_string(img) + "\n"

    # --- Step 3: Normalize lines ---
    raw_lines = [re.sub(r"\s+", " ", line.strip()) for line in text.splitlines() if line.strip()]
    merged_lines = []
    buffer = ""

    # Pattern: one or two digits + '.' + space
    serial_pattern = re.compile(r"^\s*\d{1,2}\.\s+")
    # Stop merging once tables start
    stop_keywords = ("DETAILS", "DAILY", "WORKING", "DATE", "HOURS")

    merging_allowed = True

    for line in raw_lines:
        clean = line.strip()
        if not clean:
            continue

        # --- Stop condition ---
        if any(clean.upper().startswith(word) for word in stop_keywords):
            if buffer:
                merged_lines.append(buffer.strip())
                buffer = ""
            merging_allowed = False

        if merging_allowed:
            if serial_pattern.match(clean):
                # New numbered entry → flush buffer
                if buffer:
                    merged_lines.append(buffer.strip())
                    buffer = ""
                buffer = clean
            else:
                # Continuation line
                if buffer:
                    buffer += " " + clean
                else:
                    merged_lines.append(clean)
        else:
            if buffer:
                merged_lines.append(buffer.strip())
                buffer = ""
            merged_lines.append(clean)

    # Final flush
    if buffer:
        merged_lines.append(buffer.strip())

    # --- Step 4: Post-cleaning ---
    cleaned_output = []
    for line in merged_lines:
        line = re.sub(r"\s{2,}", " ", line)  # collapse spaces
        line = line.replace("’", "'").replace("`", "'")
        cleaned_output.append(line)

    return cleaned_output, method


def extract_docx_lines(path: Path) -> List[str]:
    return [p.text.strip() for p in Document(path).paragraphs if p.text.strip()]


def extract_txt_lines(path: Path) -> List[str]:
    with open(path, "r", encoding="utf-8") as f:
        return [line.strip() for line in f if line.strip()]


# ---------------- Cleaning (file1.py logic) ---------------- #
def clean_and_merge(lines: List[str]) -> List[str]:
    serial_pattern = re.compile(r"^\s*\d+\.\s+")
    stop_keywords = ("DETAILS", "OF", "DAILY", "WORKING", "Date")

    merged_lines = []
    buffer = ""
    merging_allowed = True

    for line in lines:
        clean = line.strip()
        if not clean:
            continue

        if any(clean.upper().startswith(word.upper()) for word in stop_keywords):
            if buffer:
                merged_lines.append(buffer.strip())
                buffer = ""
            merging_allowed = False

        if merging_allowed:
            if serial_pattern.match(clean):
                if buffer:
                    merged_lines.append(buffer.strip())
                    buffer = ""
                buffer = clean
            else:
                if buffer:
                    buffer += " " + clean
                else:
                    merged_lines.append(clean)
        else:
            if buffer:
                merged_lines.append(buffer.strip())
                buffer = ""
            merged_lines.append(clean)

    if buffer:
        merged_lines.append(buffer.strip())

    return merged_lines


# ---------------- Parsing (file2.py logic) ---------------- #
time_pattern = re.compile(r"\b\d{4}\b")
range_pattern = re.compile(r"\b(\d{4})\s*[-|]\s*(\d{4})\b")
date_pattern = re.compile(r"\b\d{2}[./-]\d{2}[./-]\d{2,4}\b", re.IGNORECASE)
day_pattern = re.compile(r"\b(MON|TUE|WED|THU|FRI|SAT|SUN)\b", re.IGNORECASE)


def normalize_time(t: str) -> Optional[str]:
    if not t:
        return None
    t = t.strip()
    if len(t) == 4 and t.isdigit():
        return f"{t[:2]}:{t[2:]}"
    return None


def normalize_date(d: str) -> Optional[str]:
    if not d:
        return None
    d = d.replace(".", "-").replace("/", "-")
    parts = d.split("-")
    if len(parts[2]) == 2:
        year = "20" + parts[2]
    else:
        year = parts[2]
    return f"{parts[0]}-{parts[1]}-{year}"


def parse_events(lines: List[str]) -> List[dict]:
    events = []
    last_date, last_day, last_end_time = None, None, None

    for raw in lines:
        date_match = date_pattern.search(raw)
        day_match = day_pattern.search(raw)
        if date_match:
            last_date = normalize_date(date_match.group(0))
        if day_match:
            last_day = day_match.group(0).capitalize()

        range_match = range_pattern.search(raw)
        times = time_pattern.findall(raw)

        start_time, end_time = None, None
        if range_match:
            start_time = normalize_time(range_match.group(1))
            end_time = normalize_time(range_match.group(2))
        elif len(times) == 1:
            end_time = normalize_time(times[0])
            start_time = last_end_time if last_end_time else "00:00"
        elif len(times) >= 2:
            start_time = normalize_time(times[0])
            end_time = normalize_time(times[1])

        if start_time or end_time:
            event_name = raw
            for t in times:
                event_name = event_name.replace(t, "")
            if date_match:
                event_name = event_name.replace(date_match.group(0), "")
            if day_match:
                event_name = event_name.replace(day_match.group(0), "")
            event_name = re.sub(r"\bhrs?\b", "", event_name, flags=re.IGNORECASE)
            event_name = re.sub(r"[-|]", " ", event_name)
            event_name = re.sub(r"\s+", " ", event_name).strip()

            # eventName NULL if it has no alphabets
            if not re.search(r"[A-Za-z]", event_name):
                event_name = "NULL"

            events.append({
                "s_no": len(events) + 1,
                "eventName": event_name,
                "date": f"{last_date or ''} {last_day or ''}".strip() or "NULL",
                "startTime": start_time or "NULL",
                "endTime": end_time or "NULL",
            })
            last_end_time = end_time

    return events


# ---------------- Document Processor ---------------- #
def process_documents_with_ml(sof_file_path: Path) -> List[dict]:
    ext = sof_file_path.suffix.lower().lstrip(".")
    if ext == "pdf":
        sof_lines, _ = pdf_extract(sof_file_path)
    elif ext == "docx":
        sof_lines = extract_docx_lines(sof_file_path)
    elif ext == "txt":
        sof_lines = extract_txt_lines(sof_file_path)
    else:
        raise HTTPException(400, "Unsupported file type")

    cleaned = clean_and_merge(sof_lines)
    events = parse_events(cleaned)

    if events:
        try:
            pd.DataFrame(events).to_csv(Path(settings.UPLOAD_DIR) / "combined_output.csv", index=False)
            with open(Path(settings.UPLOAD_DIR) / "combined_output.json", "w") as jf:
                json.dump(events, jf, indent=4)
        except Exception as e:
            print(f"Error saving debug output: {e}")

    return events


# ---------------- API ---------------- #
@app.post("/process-documents/", status_code=status.HTTP_202_ACCEPTED)
async def process_documents(sof_file: UploadFile = File(...)):
    allowed_ext = ["pdf", "docx", "txt"]
    file_extension = Path(sof_file.filename).suffix.lower().lstrip(".")
    if file_extension not in allowed_ext:
        raise HTTPException(400, f"Invalid file type. Allowed: {', '.join(allowed_ext)}")

    sof_file_name = f"{uuid.uuid4()}_{sof_file.filename}"
    sof_file_path = Path(settings.UPLOAD_DIR) / sof_file_name
    with sof_file_path.open("wb") as buffer:
        shutil.copyfileobj(sof_file.file, buffer)

    processed_data = process_documents_with_ml(sof_file_path)

    return JSONResponse(
        status_code=200,
        content={
            "message": "Files processed successfully!",
            "processed_data": processed_data,
            "csv_path": str(Path(settings.UPLOAD_DIR) / "combined_output.csv"),
            "json_path": str(Path(settings.UPLOAD_DIR) / "combined_output.json"),
        },
    )


@app.get("/")
async def root():
    return {"message": "Welcome to the Document Processor Backend with ML + NLP!"}
